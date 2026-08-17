"""
SECUROXI AI DOCX Document Parser
Extracts text paragraphs, tables, font attributes, and hidden (w:vanish) text properties.
"""

import os
import zipfile
import xml.etree.ElementTree as ET
from typing import List, Optional
from securoxi.config import SecuroxiConfig
from securoxi.logger import get_logger
from securoxi.models import TextSpan
from securoxi.parsers.base import BaseParser

# Try importing python-docx if installed
HAS_DOCX = False
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DOCXParser(BaseParser):
    """
    Production-grade DOCX parser for SECUROXI AI.
    Extracts text spans, font size, hex colors, and w:vanish / w:hidden text flags.
    """

    def __init__(self, config: Optional[SecuroxiConfig] = None):
        self.config = config or SecuroxiConfig()
        self.logger = get_logger("securoxi.parsers.docx")

    def parse(self, file_path: str) -> List[TextSpan]:
        canonical_path = os.path.abspath(os.path.realpath(file_path))
        if not os.path.isfile(canonical_path):
            raise FileNotFoundError(f"DOCX file not found: '{file_path}'")

        file_size = os.path.getsize(canonical_path)
        if file_size > self.config.max_file_size_bytes:
            raise ValueError(f"DOCX file size ({file_size} bytes) exceeds maximum limit.")

        spans: List[TextSpan] = []

        if HAS_DOCX:
            try:
                doc = docx.Document(canonical_path)
                for p_idx, p in enumerate(doc.paragraphs):
                    for r_idx, r in enumerate(p.runs):
                        text = r.text
                        if not text or not text.strip():
                            continue

                        font_name = r.font.name or "Calibri"
                        font_size = float(r.font.size.pt) if r.font.size else 11.0
                        
                        # Extract color
                        font_color = "#000000"
                        if r.font.color and r.font.color.rgb:
                            font_color = f"#{r.font.color.rgb}"

                        # Detect hidden text flag (w:vanish)
                        is_hidden = bool(r.font.hidden)

                        span = TextSpan(
                            text=text,
                            page=1,
                            font_name=font_name,
                            font_size=font_size,
                            font_color=font_color,
                            bg_color="#FFFFFF",
                            is_hidden=is_hidden,
                            source="NATIVE_DOCX",
                            metadata={
                                "source": "NATIVE_DOCX",
                                "paragraph_index": p_idx,
                                "run_index": r_idx,
                                "is_hidden": is_hidden
                            }
                        )
                        spans.append(span)

                # Process Tables
                for t_idx, table in enumerate(doc.tables):
                    for row in table.rows:
                        for cell in row.cells:
                            text = cell.text.strip()
                            if text:
                                span = TextSpan(
                                    text=text,
                                    page=1,
                                    font_name="Calibri",
                                    font_size=10.0,
                                    font_color="#000000",
                                    bg_color="#FFFFFF",
                                    source="NATIVE_DOCX",
                                    metadata={"source": "NATIVE_DOCX", "table_index": t_idx}
                                )
                                spans.append(span)

                if spans:
                    return spans
            except Exception as err:
                self.logger.warning(f"python-docx parsing warning: {err}. Using XML fallback parser.")

        # Fallback XML parsing over word/document.xml inside DOCX zip container
        try:
            with zipfile.ZipFile(canonical_path, "r") as z:
                xml_content = z.read("word/document.xml")
                root = ET.fromstring(xml_content)
                namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

                for p in root.findall(".//w:p", namespaces):
                    for r in p.findall(".//w:r", namespaces):
                        t_elem = r.find("w:t", namespaces)
                        if t_elem is not None and t_elem.text:
                            text = t_elem.text
                            if not text.strip():
                                continue

                            rPr = r.find("w:rPr", namespaces)
                            is_hidden = False
                            if rPr is not None and (rPr.find("w:vanish", namespaces) is not None or rPr.find("w:hidden", namespaces) is not None):
                                is_hidden = True

                            span = TextSpan(
                                text=text,
                                page=1,
                                font_name="Calibri",
                                font_size=11.0,
                                font_color="#000000",
                                bg_color="#FFFFFF",
                                is_hidden=is_hidden,
                                source="NATIVE_DOCX",
                                metadata={"source": "NATIVE_DOCX", "is_hidden": is_hidden}
                            )
                            spans.append(span)
        except Exception as xml_err:
            self.logger.error(f"DOCX XML fallback parsing failed: {xml_err}")
            raise ValueError(f"Failed to parse DOCX document: {xml_err}")

        return spans
